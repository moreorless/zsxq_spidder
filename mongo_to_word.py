# -*- coding: utf-8 -*-
# mongodb to word
import datetime
import re

from pymongo import MongoClient
from docx import Document
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from urllib.parse import quote
from urllib.parse import unquote


# 连接mongodb的test数据库
mongo_url = "mongodb://127.0.0.1:27017/?directConnection=true&serverSelectionTimeoutMS=2000"
client = MongoClient(mongo_url)
coll = client.zsxq.gongzhou

# print("find all posts:")
# topics = coll.find({"type": "q&a"})
# for topic in topics:
#     # print(topic)
#     pprint(topic["question"]["text"])


# 按年、分类导出文章
def find_by_year_and_type(year, topicType):
    return coll.find({
        "create_time": {'$regex': year + '.*'},
        "type": topicType
    })


# topics = find_by_year_and_type('2021', 'q&a')
# for topic in topics:
#     pprint(topic["question"]["text"])

# 过滤不可见字符
def filter_non_printable_char(str):
    for i in range(0, 32):
        if i == 10 or i == 13:
            continue
        str = str.replace(chr(i), '')
    str = str.replace(chr(127), '')
    return str


def write_qa(name, topics):
    # topics.sort(key=lambda x: x.get('id'))
    document = Document()
    document.add_heading(name)

    style = document.styles['Normal']
    style.font.name = 'Times New Roman'   # 必须先设置font.name
    # style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    num = 1
    for topic in topics:
        create_time = datetime.datetime.strptime(topic["create_time"], "%Y-%m-%dT%H:%M:%S.%f+0800")\
            .strftime("%Y-%m-%d %H:%M:%S")
        likes_count = topic["likes_count"]

        # 写入标题
        para_heading = document.add_heading('', level=2)

        title = "问题"
        if topic["digested"]:
            title = "精华问题"
        title = title + str(num) + "    " + create_time + "    点赞数: " + str(likes_count)
        para_heading.add_run(title)

        # 写入问题和答案
        question = handle_text(topic["question"]["text"])
        answer = handle_text(topic["answer"]["text"])
        paragraph = document.add_paragraph()
        paragraph.add_run(u"{}".format(question)).bold = True
        try:
            document.add_paragraph(u"{}".format(answer))
        except:
            print("topic id " + str(topic["topic_id"]) + "---> " + answer)
        document.add_paragraph()

        num += 1

    export_file = "export/" + name + ".docx";
    document.save(export_file)
    print("导出文件: " + export_file)


def write_talk(name, topics):
    # topics.sort(key=lambda x: x.get('id'))
    document = Document()
    document.add_heading(name)

    style = document.styles['Normal']
    style.font.name = 'Times New Roman'   # 必须先设置font.name
    # style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    num = 1
    for topic in topics:
        create_time = datetime.datetime.strptime(topic["create_time"], "%Y-%m-%dT%H:%M:%S.%f+0800") \
            .strftime("%Y-%m-%d %H:%M:%S")
        likes_count = topic["likes_count"]

        # 写入标题
        para_heading = document.add_heading('', level=2)

        title = ""
        if topic["digested"]:
            title = "精华"
        title = title + "talk " + str(num) + "    " + create_time + "    点赞数: " + str(likes_count)
        para_heading.add_run(title)

        # 写入问题和答案
        talk_text = handle_text(topic["talk"]["text"])
        paragraph = document.add_paragraph()
        paragraph.add_run(u"{}".format(talk_text))
        document.add_paragraph()
        num += 1

    export_file = "export/" + name + ".docx";
    document.save(export_file)
    print("导出文件: " + export_file)


def handle_text(text):
    return handle_link(filter_non_printable_char(text))


def handle_link(text):
    soup = BeautifulSoup(text, "html.parser")

    mention = soup.find_all('e', attrs={'type': 'mention'})
    if len(mention):
        for m in mention:
            mention_name = m.attrs['title']
            new_tag = soup.new_tag('span')
            new_tag.string = mention_name
            m.replace_with(new_tag)

    hashtag = soup.find_all('e', attrs={'type': 'hashtag'})
    if len(hashtag):
        for tag in hashtag:
            tag_name = unquote(tag.attrs['title'])
            new_tag = soup.new_tag('span')
            new_tag.string = tag_name
            tag.replace_with(new_tag)

    links = soup.find_all('e', attrs={'type': 'web'})
    if len(links):
        for link in links:
            title = unquote(link.attrs['title'])
            href = unquote(link.attrs['href'])
            # new_a_tag = soup.new_tag('a', href=href)
            # new_a_tag.string = title
            # link.replace_with(new_a_tag)
            link.replace_with(href)

    text = str(soup)
    text = re.sub(r'<e[^>]*>', '', text).strip()
    return text


# for year in range(2019, 2022):
#     topics = find_by_year_and_type(str(year), 'q&a')
#     write_qa("问答-" + str(year), topics)

for year in range(2019, 2022):
    topics = find_by_year_and_type(str(year), 'talk')
    write_talk("讲解-" + str(year), topics)
